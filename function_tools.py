"""
================================================================================
tool函数库 (function_tools.py)

提供各种基础功能：
1. 向量数据库操作（存数据、搜数据）
2. 文档读取
3. AI模型调用
4. 工具函数（中文转拼音等）
================================================================================
"""

# 导入需要的库
import os  # 用于读取环境变量
import uuid  # 生成唯一ID的工具

import chromadb  # ChromaDB向量数据库
from chromadb.config import Settings  # ChromaDB配置
from models import *  # 从models.py导入所有模型配置
from functools import wraps  # 装饰器工具
from pypinyin import pinyin, Style  # 中文转拼音
from docx import Document as DocxDocument  # 读取Word文档
from langchain_text_splitters import RecursiveCharacterTextSplitter  # 文本分割工具
from langchain_chroma import Chroma as LangChainChroma  # LangChain向量存储
from langchain_core.documents import Document  # LangChain文档对象
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import MarkdownListOutputParser

from logger import logger


# =============================================================================
# 第一部分：向量数据库类
# =============================================================================

class MyVectorDBConnector:
    """
    【类名】MyVectorDBConnector
    
    【作用】封装ChromaDB向量数据库的操作
    
    【核心概念】
    - Collection（集合）：类似于关系数据库中的"表"，每个文档一个集合
    - Document（文档）：存储的文本内容
    - Embedding（向量）：文本的数字表示
    - Query（查询）：搜索相似向量的过程
    """
    
    def __init__(self):
        """
        【构造函数】创建向量数据库连接对象时自动调用
        
        【作用】初始化数据库连接
        """
        # 创建ChromaDB客户端
        # PersistentClient表示"持久化客户端"，数据会保存到本地文件
        # path="./chroma" 表示数据保存在当前目录的chroma文件夹中
        self.chroma_client = chromadb.PersistentClient(path="./chroma")
        
        # 创建AI模型客户端（用于生成向量）
        self.client = get_huayan_model_client()
        
        logger.info("【向量数据库】连接成功，数据保存在 ./chroma 文件夹")

    def search(self, query, collection_name='demo', n_results=10, hybrid_search=False, vector_weight=0.5):
        """
        【功能】在向量数据库中搜索相关内容
        
        【参数】
            query: 查询文本（用户的问题）
            collection_name: 在哪个集合中搜索
            n_results: 返回最相关的几条结果
            hybrid_search: 是否启用混合检索（向量检索 + 关键词检索）
            vector_weight: 向量检索的权重（0~1），关键词检索权重 = 1 - vector_weight
        
        【返回值】
            字典，包含：
            - documents: 检索到的文档内容列表
            - ids: 文档ID列表
            - distances: 相似度距离列表（越小越相似）
            - metadatas: 元数据列表
        """
        logger.debug(f"hybrid_search: {hybrid_search}")
        if hybrid_search:
            # 混合检索模式：向量检索 + 关键词检索
            try:
                from langchain_classic.retrievers import EnsembleRetriever
                from langchain_community.retrievers import BM25Retriever

                # 获取集合中的所有文档文本（用于BM25关键词检索）
                collection = self.chroma_client.get_collection(name=collection_name)
                all_docs = collection.get()
                doc_texts = all_docs['documents']

                if not doc_texts:
                    logger.warning('【混合检索警告】集合为空，使用标准向量检索')
                else:
                    logger.info(f'【混合检索】正在搜索: "{query}"')
                    logger.info(f'关键词检索: "{1 - vector_weight}", 向量检索: "{vector_weight}"')
                    logger.info(f'【混合检索】集合: {collection_name}, 返回{n_results}条结果')
                    # 创建向量检索器
                    embeddings = get_huayan_embeddings()
                    vectorstore = LangChainChroma(
                        client=self.chroma_client,
                        collection_name=collection_name,
                        embedding_function=embeddings
                    )
                    vector_retriever = vectorstore.as_retriever(search_kwargs={"k": n_results})

                    # 创建关键词检索器
                    bm25_retriever = BM25Retriever.from_texts(doc_texts)
                    bm25_retriever.k = n_results

                    # 组合成混合检索器（参考test.py的实现）
                    keyword_weight = 1 - vector_weight
                    ensemble_retriever = EnsembleRetriever(
                        retrievers=[bm25_retriever, vector_retriever],
                        weights=[keyword_weight, vector_weight]
                    )

                    # 执行混合检索
                    docs = ensemble_retriever.invoke(query)

                    logger.info(f'【混合检索】完成，找到 {len(docs)} 条结果')

                    return {
                        'documents': [[doc.page_content for doc in docs]],
                        'ids': [[str(uuid.uuid4()) for _ in docs]],
                        'distances': [[0.0] * len(docs)],
                        'metadatas': [[doc.metadata for doc in docs]]
                    }
            except ImportError as e:
                logger.error(f'【混合检索错误】缺少依赖库: {e}，回退到标准向量检索')
            except Exception as e:
                logger.error(f'【混合检索错误】{e}，回退到标准向量检索')
        
        # 使用 LangChain 检索器进行向量相似度搜索
        logger.info(f'【向量数据库】正在搜索: "{query}"')
        logger.info(f'【向量数据库】集合: {collection_name}, 返回{n_results}条结果')
        embeddings = get_huayan_embeddings()
        vectorstore = LangChainChroma(
            client=self.chroma_client,
            collection_name=collection_name,
            embedding_function=embeddings
        )

        retriever = vectorstore.as_retriever(
            search_type="similarity",
            search_kwargs={"k": n_results}
        )

        docs = retriever.invoke(query)

        logger.info(f'【向量数据库】搜索完成，找到 {len(docs)} 条结果')

        return {
            'documents': [[doc.page_content for doc in docs]],
            'ids': [[str(uuid.uuid4()) for _ in docs]],
            'distances': [[0.0] * len(docs)],
            'metadatas': [[doc.metadata for doc in docs]]
        }

    def delete_collection(self, collection_name):
        """
        【功能】删除向量数据库中的集合
        
        【参数】
            collection_name: 要删除的集合名称
        
        【返回值】
            bool: 删除是否成功
        """
        try:
            logger.info(f'【向量数据库】正在删除集合: {collection_name}')
            # 调用ChromaDB客户端的delete_collection方法删除集合
            self.chroma_client.delete_collection(name=collection_name)
            logger.info(f'【向量数据库】集合 {collection_name} 删除成功')
            return True
        except Exception as e:
            logger.error(f'【向量数据库】删除集合 {collection_name} 失败: {str(e)}')
            return False

# =============================================================================
# 第二部分：AI模型调用函数
# =============================================================================

def get_completion(info, user_query, history="", model="qwen3.5-flash"):
    """
    【功能】调用大语言模型获取回答
    
    【参数】
        info: 已知信息（从知识库检索到的内容）
        user_query: 用户的问题
        history: 历史对话记录
        model: 使用的模型名称
    
    【返回值】
        AI生成的文本回答
    """

    # 1. 创建提示词模板
    prompt = PromptTemplate(template="""你是一个专业的企业知识库问答助手。
    你的任务是根据下述给定的已知信息和历史对话回答用户问题。
    请确保你的回复完全依据下述已知信息，不要编造答案。

    【重要规则】
    1. 如果用户询问的指标在已知信息中直接存在，请直接返回该数值
    2. 如果用户询问的指标在已知信息中不存在，但可以通过计算得出（如总计=各项之和、均值=总和/数量等），请进行简单计算后回答，但不要自己推理
    3. 如果无法从已知信息中获取或计算出答案，才回复"根据现有资料，我无法回答您的问题"
    4. 在回答时，如果进行了简单的计算，除了给出最终结果，最好简要说明计算过程
    5. 请参考历史对话，保持回答的连贯性
    6. 务必以Markdown格式返回结果

    【历史对话】:
    {history}

    【已知信息】:
    {info}

    ================================================================================
    【用户问题】：
    {user_query}
    ================================================================================

    请用中文回答用户问题，回答要简洁明了。

    【输出格式要求】
    请使用 Markdown 列表格式输出你的回答，每条内容前使用 "- " 开头。""")

    # 2. 获取AI客户端（使用LangChain版本）
    client = get_huayan_model_client()

    # 3. 构建chain（不使用MarkdownListOutputParser，直接返回字符串）
    chain = prompt | client

    # 4. 使用invoke方法调用chain
    # 传递输入参数：info、user_query和history
    response = chain.invoke({
        "info": info,
        "user_query": user_query,
        "history": history
    })
    
    # 5. 返回结果（确保返回的是字符串）
    if hasattr(response, 'content'):
        return response.content
    return response


# =============================================================================
# 第四部分：工具装饰器
# =============================================================================

def to_pinyin(fn):
    """
    【功能】装饰器：将中文集合名转换为拼音
    
    【作用】
    有些数据库不支持中文作为标识符，所以把中文转成拼音更安全。
    """
    @wraps(fn)  # 保留原函数的元信息
    def chinese_to_pinyin(*args, **kwargs):
        # 获取collection_name参数
        chinese_name = kwargs['collection_name']
        
        # 去掉文件名中的点号（.）
        chinese_name = chinese_name.replace('.', '')
        
        # 使用pypinyin库将中文转为拼音
        # style=Style.NORMAL: 普通风格，不带声调
        # heteronym=False: 不使用多音字
        pinyin_list = pinyin(chinese_name, style=Style.NORMAL, heteronym=False)
        
        # 将拼音列表拼接成字符串
        # word[0] 取每个字的第一个拼音
        # .lower() 转成小写
        pinyin_str = ''.join([word[0].lower() for word in pinyin_list])
        
        # 替换参数中的中文名为拼音
        kwargs['collection_name'] = pinyin_str
        
        # 调用原函数
        return fn(*args, **kwargs)
    
    return chinese_to_pinyin


def rerank_documents(docs, query, top_n=5):
    """
    【功能】对检索结果进行重排序
    
    【参数】
        docs: 文档内容列表
        query: 用户查询
        top_n: 保留前N条结果
    
    【返回值】
        重排序后的文档内容列表
    """
    if not get_huayan_rerank or not docs:
        return docs
    
    try:
        logger.info('\n>>> 执行重排序...')
        reranker = get_huayan_rerank(top_n=top_n)
        
        # 构建文档对象列表
        from langchain_core.documents import Document
        doc_objects = [Document(page_content=doc) for doc in docs]
        
        # 执行重排序
        reranked_docs = reranker.compress_documents(doc_objects, query=query)
        
        # 提取重排序后的文档内容
        result = [doc.page_content for doc in reranked_docs]
        
        logger.info(f'重排序完成，保留前 {len(result)} 条结果')
        return result
    except Exception as e:
        logger.error(f'【重排序错误】{e}，返回原始结果')
        return docs


def check_model_health(model_type, test_type='light'):
    """
    【功能】检查模型健康状态
    
    【参数】
        model_type: 模型类型 ('chat', 'embedding', 'rerank')
        test_type: 检测类型 ('light', 'full')
    
    【返回值】
        dict: 模型健康状态信息
    """
    import time
    from datetime import datetime
    
    start_time = time.time()
    result = {
        'model': '',
        'status': 'unknown',
        'message': '',
        'response_time': 0,
        'last_check': datetime.now().isoformat(),
        'api_key_configured': False
    }
    
    try:
        # 检查API密钥
        if model_type in ['chat', 'embedding', 'rerank']:
            api_key = API_KEY
            result['api_key_configured'] = bool(api_key)
            if not api_key:
                result['status'] = 'offline'
                result['message'] = 'API密钥未配置'
                result['response_time'] = int((time.time() - start_time) * 1000)
                return result
        
        # 轻量级检测
        if test_type == 'light':
            if model_type == 'chat':
                result['model'] = HUAYAN_CODE_REASONING_MODEL
                # 尝试初始化客户端
                from models import get_huayan_model_client
                client = get_huayan_model_client()
                result['status'] = 'online'
                result['message'] = '模型客户端初始化成功'
                
            elif model_type == 'embedding':
                result['model'] = HUAYAN_EMBEDDING_MODEL
                # 尝试初始化嵌入模型
                from models import get_huayan_embeddings
                embeddings = get_huayan_embeddings()
                result['status'] = 'online'
                result['message'] = '嵌入模型初始化成功'
                
            elif model_type == 'rerank':
                result['model'] = HUAYAN_RERANK_MODEL
                # 尝试初始化重排序模型
                from models import get_huayan_rerank
                reranker = get_huayan_rerank(top_n=3)
                result['status'] = 'online'
                result['message'] = '重排序模型初始化成功'
                
        # 完整检测
        elif test_type == 'full':
            if model_type == 'chat':
                result['model'] = HUAYAN_CODE_REASONING_MODEL
                from models import get_huayan_model_client
                client = get_huayan_model_client()
                # 发送测试提示
                test_prompt = "测试消息"
                from langchain_core.messages import HumanMessage
                response = client.invoke([HumanMessage(content=test_prompt)])
                if response.content:
                    result['status'] = 'online'
                    result['message'] = '对话模型测试成功'
                
            elif model_type == 'embedding':
                result['model'] = HUAYAN_EMBEDDING_MODEL
                from models import get_huayan_embeddings
                embeddings = get_huayan_embeddings()
                # 生成测试向量
                test_text = "测试文本"
                embedding = embeddings.embed_query(test_text)
                if embedding and len(embedding) > 0:
                    result['status'] = 'online'
                    result['message'] = '嵌入模型测试成功'
                
            elif model_type == 'rerank':
                result['model'] = HUAYAN_RERANK_MODEL
                from models import get_huayan_rerank
                from langchain_core.documents import Document
                reranker = get_huayan_rerank(top_n=3)
                # 测试重排序
                test_docs = [
                    Document(page_content="这是测试文档1"),
                    Document(page_content="这是测试文档2"),
                    Document(page_content="这是测试文档3")
                ]
                test_query = "测试查询"
                reranked = reranker.compress_documents(test_docs, query=test_query)
                if reranked:
                    result['status'] = 'online'
                    result['message'] = '重排序模型测试成功'
        
    except Exception as e:
        result['status'] = 'error'
        result['message'] = f'检测失败: {str(e)}'
    
    result['response_time'] = int((time.time() - start_time) * 1000)
    return result


def check_mineru_health():
    """
    【功能】检查MinerU PDF解析引擎是否可用
    
    【返回值】
        dict: MinerU健康状态信息
    """
    import time
    from datetime import datetime
    
    start_time = time.time()
    result = {
        'model': 'MinerU',
        'status': 'unknown',
        'message': '',
        'response_time': 0,
        'last_check': datetime.now().isoformat(),
        'api_key_configured': True
    }
    
    try:
        try:
            from mineru.cli.common import do_parse, read_fn
            result['status'] = 'online'
            result['message'] = 'MinerU解析引擎可用'
        except ImportError:
            result['status'] = 'offline'
            result['message'] = 'MinerU未安装，仅支持PyPDF2基础解析'
    except Exception as e:
        result['status'] = 'error'
        result['message'] = f'检测失败: {str(e)}'
    
    result['response_time'] = int((time.time() - start_time) * 1000)
    return result


def check_all_models_health(test_type='light'):
    """
    【功能】检查所有模型的健康状态
    
    【参数】
        test_type: 检测类型 ('light', 'full')
    
    【返回值】
        dict: 所有模型的健康状态信息
    """
    return {
        'chat': check_model_health('chat', test_type),
        'embedding': check_model_health('embedding', test_type),
        'rerank': check_model_health('rerank', test_type),
        'mineru': check_mineru_health()
    }
